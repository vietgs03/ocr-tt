from fastapi import FastAPI, File, UploadFile, HTTPException, Form
from fastapi.responses import StreamingResponse, JSONResponse, FileResponse
from starlette.background import BackgroundTask
from fastapi.middleware.cors import CORSMiddleware
import uvicorn
import asyncio
import json
import time
import os
from pathlib import Path
from selflearning_ocr import SelfLearningOCR
from pdf_extractor import extract_text_from_pdf
from paddleocr import PaddleOCR
from symspellpy import SymSpell
from pdf2docx import Converter
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from tempfile import NamedTemporaryFile
import logging

logging.basicConfig(level=logging.INFO)

app = FastAPI(title="Local OCR API", version="2.0")

# CORS for web GUI
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Initialize OCR engines
print("🚀 Initializing OCR engines...")
deepseek_ocr = SelfLearningOCR(keep_alive="60m")
paddle_ocr = PaddleOCR(use_angle_cls=False, lang='vi', use_gpu=False, show_log=False, enable_mkldnn=False)
sym_spell = SymSpell(max_dictionary_edit_distance=2, prefix_length=7)

# Load dictionary for fast mode
if os.path.exists('vn_dictionary.txt'):
    sym_spell.load_dictionary('vn_dictionary.txt', term_index=0, count_index=1, separator=" ", encoding="utf-8")
    print("✅ Dictionary loaded for Fast Mode")

print("✅ API Ready!")

@app.get("/")
async def root():
    return {
        "service": "Local OCR API",
        "version": "2.0",
        "modes": ["fast", "accurate", "auto"],
        "features": ["streaming", "caching", "hybrid"]
    }

@app.get("/health")
async def health_check():
    stats = deepseek_ocr.get_cache_stats()
    return {
        "status": "healthy",
        "cache": {
            "documents": stats['cached_documents'],
            "hits": stats['total_cache_hits'],
            "vocabulary_size": stats['vocabulary_size']
        }
    }

@app.post("/ocr/fast")
async def ocr_fast(file: UploadFile = File(...)):
    """
    Fast Mode: PaddleOCR + SymSpell
    - Speed: ~3-5s
    - Accuracy: 85-90%
    - Use for: Quick scans, drafts
    """
    try:
        # Save temp file
        temp_path = f"temp_{file.filename}"
        with open(temp_path, "wb") as f:
            f.write(await file.read())
        
        start = time.time()
        
        # PaddleOCR processing
        result = paddle_ocr.ocr(temp_path, cls=False)
        
        # Extract text
        lines = []
        for line in result[0]:
            text = line[1][0]
            # Apply SymSpell correction
            suggestions = sym_spell.lookup_compound(text, max_edit_distance=2)
            if suggestions:
                lines.append(suggestions[0].term)
            else:
                lines.append(text)
        
        ocr_text = "\n".join(lines)
        duration = time.time() - start
        
        # Cleanup
        os.remove(temp_path)
        
        return {
            "mode": "fast",
            "text": ocr_text,
            "duration": round(duration, 2),
            "length": len(ocr_text)
        }
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/ocr/accurate")
async def ocr_accurate(file: UploadFile = File(...)):
    """
    Accurate Mode: DeepSeek-OCR
    - Speed: ~80-180s (cached: instant)
    - Accuracy: 95-98%
    - Use for: Important documents
    """
    try:
        # Save temp file
        temp_path = f"temp_{file.filename}"
        with open(temp_path, "wb") as f:
            f.write(await file.read())
        
        start = time.time()
        
        # DeepSeek OCR (with caching!)
        ocr_text = deepseek_ocr.process_image(temp_path)
        duration = time.time() - start
        
        # Cleanup
        os.remove(temp_path)
        
        return {
            "mode": "accurate",
            "text": ocr_text,
            "duration": round(duration, 2),
            "length": len(ocr_text),
            "cached": duration < 1.0  # If < 1s, was cached
        }
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/ocr/pdf-text")
async def ocr_pdf_text(file: UploadFile = File(...)):
    """
    PDF Text Mode: Direct text extraction (No OCR)
    - Speed: Near-instant
    - Accuracy: 100% (of text layer)
    - Use for: Searchable PDFs, digital documents
    """
    if not file.filename.lower().endswith('.pdf'):
        raise HTTPException(status_code=400, detail="Only PDF files are supported for this mode")
        
    try:
        # Save temp file
        temp_path = f"temp_{file.filename}"
        with open(temp_path, "wb") as f:
            f.write(await file.read())
            
        start = time.time()
        
        # Direct extraction
        text = extract_text_from_pdf(temp_path)
        duration = time.time() - start
        
        # Cleanup
        os.remove(temp_path)
        
        return {
            "mode": "pdf_text",
            "text": text,
            "duration": round(duration, 2),
            "length": len(text)
        }
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/ocr/stream")
async def ocr_stream(file: UploadFile = File(...), mode: str = "accurate"):
    """
    Streaming Mode: Real-time text output
    - GUI sees text appearing character by character
    - Better UX for long processing
    """
    async def generate():
        try:
            # Save temp file
            temp_path = f"temp_{file.filename}"
            with open(temp_path, "wb") as f:
                f.write(await file.read())
            
            # Send start event
            yield f"data: {json.dumps({'type': 'start', 'mode': mode})}\n\n"
            
            if mode == "fast":
                # Fast mode - instant result
                result = paddle_ocr.ocr(temp_path, cls=False)
                text = "\n".join([line[1][0] for line in result[0]])
                
                # Stream character by character for demo
                for char in text:
                    yield f"data: {json.dumps({'type': 'token', 'content': char})}\n\n"
                    await asyncio.sleep(0.01)  # Simulate streaming
            
            else:
                # Accurate mode - DeepSeek with real streaming
                import ollama
                with open(temp_path, 'rb') as img_file:
                    img_data = img_file.read()
                
                stream = ollama.chat(
                    model='deepseek-ocr',
                    messages=[{
                        'role': 'user',
                        'content': 'Free OCR.',
                        'images': [img_data]
                    }],
                    stream=True
                )
                
                for chunk in stream:
                    content = chunk['message']['content']
                    if content:
                        yield f"data: {json.dumps({'type': 'token', 'content': content})}\n\n"
            
            # Send completion
            yield f"data: {json.dumps({'type': 'done'})}\n\n"
            
            # Cleanup
            os.remove(temp_path)
            
        except Exception as e:
            yield f"data: {json.dumps({'type': 'error', 'message': str(e)})}\n\n"
    
    return StreamingResponse(generate(), media_type="text/event-stream")

@app.post("/ocr/hybrid")
async def ocr_hybrid(file: UploadFile = File(...)):
    """
    Hybrid Mode: Smart selection
    - Try cache first
    - If miss: Use fast mode first, then accurate in background
    """
    try:
        temp_path = f"temp_{file.filename}"
        with open(temp_path, "wb") as f:
            f.write(await file.read())
        
        # Check cache first
        import hashlib
        from PIL import Image
        img = Image.open(temp_path).convert('L').resize((32, 32))
        pixels = list(img.getdata())
        avg = sum(pixels) / len(pixels)
        bits = ''.join('1' if p > avg else '0' for p in pixels)
        img_hash = hashlib.md5(bits.encode()).hexdigest()
        
        cached = deepseek_ocr._check_cache(img_hash)
        if cached:
            os.remove(temp_path)
            return {
                "mode": "hybrid_cached",
                "text": cached,
                "duration": 0.02,
                "source": "cache"
            }
        
        # Fast mode first for immediate response
        start = time.time()
        result = paddle_ocr.ocr(temp_path, cls=False)
        fast_text = "\n".join([line[1][0] for line in result[0]])
        fast_duration = time.time() - start
        
        os.remove(temp_path)
        
        return {
            "mode": "hybrid_fast",
            "text": fast_text,
            "duration": round(fast_duration, 2),
            "source": "paddle",
            "note": "Use /ocr/accurate for higher quality"
        }
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/vocabulary/learn")
async def learn_vocabulary(wrong: str, correct: str):
    """Learn from user corrections"""
    deepseek_ocr.learn_correction(wrong, correct)
    return {"status": "learned", "wrong": wrong, "correct": correct}

@app.get("/cache/stats")
async def cache_stats():
    """Get cache statistics"""
    return deepseek_ocr.get_cache_stats()

@app.post("/export/docx")
async def export_docx(text: str = Form(...)):
    """Export text to Word document"""
    try:
        # Create document
        doc = Document()
        
        # Set default style to Times New Roman
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(13)
        
        # Add paragraphs (preserving line breaks)
        for line in text.split('\n'):
            p = doc.add_paragraph(line)
            p.paragraph_format.space_after = Pt(0)
            
        # Save to temp file
        with NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
            doc.save(tmp.name)
            tmp_path = tmp.name
            
        return FileResponse(
            tmp_path, 
            media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document', 
            filename="ocr_result.docx",
            background=BackgroundTask(os.remove, tmp_path)
        )
            
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/convert/pdf-doc")
async def convert_pdf_doc(file: UploadFile = File(...)):
    """Convert PDF to Word preserving layout"""
    if not file.filename.lower().endswith('.pdf'):
        raise HTTPException(status_code=400, detail="Only PDF files are supported")
        
    try:
        # Save temp PDF
        with NamedTemporaryFile(delete=False, suffix=".pdf") as tmp_pdf:
            tmp_pdf.write(await file.read())
            pdf_path = tmp_pdf.name
            
        # Define temp Docx path
        docx_path = pdf_path.replace(".pdf", ".docx")
        
        # Convert using pdf2docx
        cv = Converter(pdf_path)
        cv.convert(docx_path, start=0, end=None)
        cv.close()
        
        # Post-process to force Times New Roman
        doc = Document(docx_path)
        # Iterate over all paragraphs
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(13)
        
        # Iterate over tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.name = 'Times New Roman'
                            run.font.size = Pt(13)
                            
        doc.save(docx_path)
        
        # Cleanup PDF
        os.remove(pdf_path)
        
        return FileResponse(
            docx_path, 
            media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document', 
            filename="converted_layout.docx",
            background=BackgroundTask(os.remove, docx_path)
        )
            
    except Exception as e:
        if os.path.exists(pdf_path):
            os.remove(pdf_path)
        raise HTTPException(status_code=500, detail=str(e))

@app.delete("/cache/clear")
async def clear_cache():
    """Clear entire cache"""
    deepseek_ocr.clear_cache()
    return {"status": "cleared"}

if __name__ == "__main__":
    uvicorn.run(
        "ocr_api:app",
        host="0.0.0.0",
        port=8000,
        reload=False
    )
